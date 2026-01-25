import os
from fpdf import FPDF
from fpdf.errors import FPDFException
from docx import Document
from docx.shared import Pt
import uuid

OUTPUT_DIR = "output/generated_activities"
os.makedirs(OUTPUT_DIR, exist_ok=True)

class ActivityGenerator:
    @staticmethod
    def generate_pdf(data: dict) -> str:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        
        # Header
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "Atividade Educacional", ln=True, align="C")
        pdf.ln(10)
        
        pdf.set_font("Helvetica", size=12)
        pdf.cell(0, 10, "Escola: _______________________________________________________", ln=True)
        pdf.cell(0, 10, "Aluno: ________________________________________________________ Data: ___/___/___", ln=True)
        pdf.ln(10)
        
        # Title
        pdf.set_font("Helvetica", "B", 14)
        title = data.get("title", "Sem Título")
        safe_title = str(title).encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, safe_title, align="C")
        pdf.ln(5)
        
        # Objective
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Objetivo:", ln=True)
        pdf.set_font("Helvetica", size=12)
        objective = data.get("objective", "")
        safe_objective = str(objective).encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 10, safe_objective)
        pdf.ln(5)

        bncc_skills = data.get("bncc_skills") or []
        if bncc_skills:
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 10, "Habilidades BNCC:", ln=True)
            pdf.set_font("Helvetica", size=11)
            for skill in bncc_skills:
                # Sanitize text for standard fonts (latin-1)
                safe_skill = str(skill).encode('latin-1', 'replace').decode('latin-1')
                pdf.set_x(pdf.l_margin)
                width = pdf.w - pdf.l_margin - pdf.r_margin
                try:
                    pdf.multi_cell(width, 8, f"- {safe_skill}")
                except FPDFException:
                    pdf.multi_cell(width, 8, "-")
            pdf.ln(5)
        
        # Content (Context)
        if data.get("content"):
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 10, "Texto de Apoio:", ln=True)
            pdf.set_font("Helvetica", size=11)
            content = data.get("content", "")
            safe_content = str(content).encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 10, safe_content)
            pdf.ln(5)
        
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Questões:", ln=True)
        pdf.set_font("Helvetica", size=12)

        questions = data.get("questions", [])
        for idx, q in enumerate(questions, 1):
            if isinstance(q, dict):
                enunciado = q.get("enunciado") or q.get("question") or ""
                alternativas = q.get("alternativas") or q.get("options") or []
                safe_enunciado = str(enunciado).encode("latin-1", "replace").decode("latin-1")
                pdf.set_x(pdf.l_margin)
                width = pdf.w - pdf.l_margin - pdf.r_margin
                try:
                    pdf.multi_cell(width, 10, f"{idx}. {safe_enunciado}")
                except FPDFException:
                    # Se até o enunciado der erro, apenas pula a linha da questão
                    pdf.cell(0, 10, f"{idx}.", ln=True)
                pdf.ln(2)
                pdf.set_font("Helvetica", size=11)
                for alt in alternativas:
                    safe_alt = str(alt).encode("latin-1", "replace").decode("latin-1")
                    pdf.set_x(pdf.l_margin)
                    width = pdf.w - pdf.l_margin - pdf.r_margin
                    try:
                        pdf.multi_cell(width, 8, safe_alt)
                    except FPDFException:
                        # Se a alternativa der erro, escreve só um traço e segue
                        pdf.cell(0, 8, "-", ln=True)
                pdf.ln(8)
                pdf.set_font("Helvetica", size=12)
            else:
                safe_question = str(q).encode("latin-1", "replace").decode("latin-1")
                pdf.multi_cell(0, 10, f"{idx}. {safe_question}")
                pdf.ln(5)
                pdf.ln(20)
            
        filename = f"atividade_{uuid.uuid4()}.pdf"
        filepath = os.path.join(OUTPUT_DIR, filename)
        pdf.output(filepath)
        
        return filename

    @staticmethod
    def generate_docx(data: dict) -> str:
        doc = Document()
        
        # Header
        doc.add_heading('Atividade Educacional', 0)
        
        p = doc.add_paragraph()
        p.add_run("Escola: _______________________________________________________\n")
        p.add_run("Aluno: ________________________________________________________ Data: ___/___/___")
        
        # Title
        doc.add_heading(data.get("title", "Sem Título"), level=1)
        
        doc.add_heading('Objetivo:', level=2)
        doc.add_paragraph(data.get("objective", ""))

        bncc_skills = data.get("bncc_skills") or []
        if bncc_skills:
            doc.add_heading('Habilidades BNCC:', level=2)
            for skill in bncc_skills:
                doc.add_paragraph(f"- {skill}")
        
        # Content
        if data.get("content"):
            doc.add_heading('Texto de Apoio:', level=2)
            doc.add_paragraph(data.get("content", ""))
        
        doc.add_heading('Questões:', level=2)
        questions = data.get("questions", [])
        for idx, q in enumerate(questions, 1):
            if isinstance(q, dict):
                enunciado = q.get("enunciado") or q.get("question") or ""
                alternativas = q.get("alternativas") or q.get("options") or []
                doc.add_paragraph(f"{idx}. {enunciado}")
                for alt in alternativas:
                    doc.add_paragraph(f"{alt}")
                doc.add_paragraph("")
            else:
                doc.add_paragraph(f"{idx}. {q}")
                doc.add_paragraph("_" * 80)
                doc.add_paragraph("")
            
        filename = f"atividade_{uuid.uuid4()}.docx"
        filepath = os.path.join(OUTPUT_DIR, filename)
        doc.save(filepath)
        
        return filename
